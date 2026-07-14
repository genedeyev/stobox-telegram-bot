"""Document loaders + ingestion CLI.

Supported formats: Markdown (with YAML front-matter metadata), plain text,
HTML, PDF, and Word (.docx). Google Docs are ingested via their exported
Markdown/DOCX/HTML — export first, drop in ``docs/``.

Front-matter example (Markdown):

    ---
    title: STBU Migration Guide
    version: "2.1"
    author: Stobox Core
    date: 2025-11-02
    category: tokenomics
    product: STBU
    language: en
    visibility: public
    source_url: https://stobox.io/docs/stbu-migration
    ---
"""

from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path

import yaml

from ..logging import get_logger
from .models import DocMeta, Document

log = get_logger(__name__)

SUPPORTED = {".md", ".markdown", ".txt", ".html", ".htm", ".pdf", ".docx"}
_FRONTMATTER = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)


def _parse_date(value: object) -> date | None:
    if isinstance(value, date):
        return value
    if isinstance(value, str):
        try:
            return datetime.fromisoformat(value).date()
        except ValueError:
            return None
    return None


def _meta_from_frontmatter(fm: dict, path: Path) -> DocMeta:
    known = {
        "title", "version", "author", "date", "category",
        "product", "language", "visibility", "confidence", "source_url",
    }
    return DocMeta(
        title=str(fm.get("title") or path.stem.replace("_", " ").title()),
        source_file=str(path),
        version=str(fm["version"]) if fm.get("version") is not None else None,
        author=fm.get("author"),
        doc_date=_parse_date(fm.get("date")),
        category=fm.get("category"),
        product=fm.get("product"),
        language=fm.get("language", "en"),
        visibility=fm.get("visibility", "public"),
        confidence=float(fm.get("confidence", 1.0)),
        source_url=fm.get("source_url"),
        extra={k: v for k, v in fm.items() if k not in known},
    )


def _read_markdown(path: Path) -> tuple[str, dict]:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    fm: dict = {}
    m = _FRONTMATTER.match(raw)
    if m:
        fm = yaml.safe_load(m.group(1)) or {}
        raw = raw[m.end():]
    return raw, fm


def _read_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text("\n")


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def load_document(path: str | Path) -> Document | None:
    """Load a single file into a :class:`Document`, or ``None`` if unsupported."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        return None
    try:
        fm: dict = {}
        if ext in (".md", ".markdown", ".txt"):
            text, fm = _read_markdown(path)
        elif ext in (".html", ".htm"):
            text = _read_html(path)
        elif ext == ".pdf":
            text = _read_pdf(path)
        elif ext == ".docx":
            text = _read_docx(path)
        else:  # pragma: no cover - guarded by SUPPORTED
            return None
    except Exception as exc:  # noqa: BLE001 - one bad file shouldn't kill ingest
        log.error("ingest.read_failed", file=str(path), error=str(exc))
        return None

    text = text.strip()
    if not text:
        return None
    meta = _meta_from_frontmatter(fm, path)
    return Document(meta=meta, text=text)


def load_directory(root: str | Path) -> list[Document]:
    root = Path(root)
    docs: list[Document] = []
    if not root.exists():
        log.warning("ingest.docs_missing", path=str(root))
        return docs
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED:
            doc = load_document(path)
            if doc:
                docs.append(doc)
    log.info("ingest.loaded", count=len(docs), root=str(root))
    return docs


def cli() -> None:
    """``stobox-ingest`` entrypoint: (re)build the index from ``docs/``."""
    import argparse
    import asyncio

    from ..config import load_config
    from ..logging import configure_logging
    from .indexer import Indexer

    parser = argparse.ArgumentParser(description="Ingest Stobox docs into the index.")
    parser.add_argument("--path", default=None, help="docs directory (defaults to config)")
    parser.add_argument("--rebuild", action="store_true", help="drop and rebuild the whole index")
    args = parser.parse_args()

    configure_logging()
    config = load_config()
    docs_path = args.path or config.get("knowledge.docs_path", "docs")

    async def run() -> None:
        indexer = await Indexer.create(config)
        n = await indexer.index_directory(docs_path, rebuild=args.rebuild)
        log.info("ingest.done", chunks_indexed=n)

    asyncio.run(run())


if __name__ == "__main__":
    cli()
